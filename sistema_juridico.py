import os
import PyPDF2
from google import genai
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from dotenv import load_dotenv

class AssistenteJuridico:
    # 1. O Construtor: Prepara a IA quando a classe é chamada
    def __init__(self, api_key):
        self.cliente_ia = genai.Client(api_key=api_key)
        self.modelo = 'gemini-2.5-flash' # O modelo rápido e gratuito
        
    # 2. Operário 1: Os Olhos (Leitor de PDF)
    def extrair_texto(self, caminho_pdf):
        print(f"📖 Lendo o arquivo: {caminho_pdf}...")
        texto = ""
        try:
            with open(caminho_pdf, "rb") as arquivo:
                leitor = PyPDF2.PdfReader(arquivo)
                for pagina in leitor.pages:
                    texto += pagina.extract_text() + "\n"
            return texto
        except Exception as e:
            print(f"❌ Erro ao ler PDF: {e}")
            return None

    # 3. Operário 2: O Cérebro (IA redigindo a peça)
    def redigir_defesa(self, texto_processo):
        print("🧠 Analisando o caso e redigindo a peça...")
        prompt = f"""
        Atue como um advogado sênior especialista em litígios cíveis.
        Analise o documento anexo e GERE o esqueleto completo de uma CONTESTAÇÃO.
        
        Estruture em:
        - SÍNTESE DOS FATOS
        - PRELIMINARES
        - DO MÉRITO
        - DOS PEDIDOS
        
        Tom técnico, formal e estritamente jurídico.
        
        [DOCUMENTO]
        {texto_processo}
        """
        resposta = self.cliente_ia.models.generate_content(
            model=self.modelo,
            contents=prompt
        )
        return resposta.text

    # 4. Operário 3: As Mãos (A Impressora Word)
    def formatar_word(self, texto_peticao, nome_saida):
        print("🖨️ Formatando o documento Word...")
        doc = Document()
        
        # Cabeçalho
        cabecalho = doc.add_paragraph()
        cabecalho.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cabecalho.add_run("EXCELENTÍSSIMO SENHOR DOUTOR JUIZ DE DIREITO DA ___ VARA CÍVEL DO FORO DE ____________")
        run.bold = True
        run.font.size = Pt(12)
        run.font.name = 'Arial'
        
        doc.add_paragraph("\n")
        
        # Corpo do texto
        for linha in texto_peticao.split('\n'):
            if linha.strip():
                p = doc.add_paragraph(linha)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                for r in p.runs:
                    r.font.size = Pt(12)
                    r.font.name = 'Arial'
        
        # Assinatura
        doc.add_paragraph("\n\nTermos em que,\nPede deferimento.\n\nLocal, Data.\n\n__________________________________\nAdvogado / OAB")
        doc.save(nome_saida)
        print(f"✅ Documento finalizado e salvo como: {nome_saida}")

    # 5. O Gerente: Puxa a linha de montagem de uma vez só!
    def executar_trabalho(self, pdf_entrada, docx_saida="Peticao_Final.docx"):
        print("🚀 Iniciando o Super Estagiário Jurídico...\n")
        texto = self.extrair_texto(pdf_entrada)
        
        if texto:
            peca = self.redigir_defesa(texto)
            self.formatar_word(peca, docx_saida)
            print("\n🎉 Trabalho concluído com sucesso!")

# ==========================================
# PAINEL DE CONTROLE 
# ==========================================
if __name__ == "__main__":
    # Carrega a chave do cofre .env
    load_dotenv(override=True)
    chave_api = os.getenv("GEMINI_API_KEY")
    
    if not chave_api:
        print("❌ Chave API não encontrada no arquivo .env!")
    else:
        # 1. Ligamos o robô
        meu_estagiario = AssistenteJuridico(chave_api)
        
        # 2. Mandamos ele trabalhar com 1 único comando!
        meu_estagiario.executar_trabalho("processo_teste.pdf", "Contestacao_Oficial.docx")