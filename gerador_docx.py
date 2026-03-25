from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def criar_peticao_word(texto_peticao, nome_saida="Peticao_Final_Pronta.docx"):
    print("🖨️ Formatando o documento para o Microsoft Word...")
    
    # Cria um documento em branco
    doc = Document()
    
    # 1. Criando o Cabeçalho Padrão (Centralizado e em Negrito)
    cabecalho = doc.add_paragraph()
    cabecalho.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_cabecalho = cabecalho.add_run("EXCELENTÍSSIMO SENHOR DOUTOR JUIZ DE DIREITO DA ___ VARA CÍVEL DO FORO DE ____________")
    run_cabecalho.bold = True
    run_cabecalho.font.size = Pt(12)
    run_cabecalho.font.name = 'Arial'
    
    # Dá uns espaços em branco
    doc.add_paragraph("\n\n")
    
    # 2. Adicionando o texto gerado pela IA (Parágrafo por parágrafo)
    paragrafos = texto_peticao.split('\n')
    
    for linha in paragrafos:
        if linha.strip(): # Se a linha não estiver vazia
            p = doc.add_paragraph(linha)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY # Texto justificado
            
            # Formatação da fonte do corpo do texto
            for run in p.runs:
                run.font.size = Pt(12)
                run.font.name = 'Arial'
    
    # 3. Assinatura no final
    doc.add_paragraph("\n\nTermos em que,\nPede deferimento.\n\nLocal, Data.\n\n__________________________________\nAdvogado / OAB")

    # Salva o arquivo no seu computador
    doc.save(nome_saida)
    print(f"✅ Documento Word gerado com sucesso! Arquivo salvo como: {nome_saida}")

# ==========================================
# PAINEL DE CONTROLE (Teste da Impressora)
# ==========================================
if __name__ == "__main__":
    try:
        # Lê o texto que a IA gerou no passo anterior
        with open("minuta_defesa.txt", "r", encoding="utf-8") as arquivo:
            texto_ia = arquivo.read()
            
        criar_peticao_word(texto_ia)
    except FileNotFoundError:
        print("❌ Erro: O arquivo 'minuta_defesa.txt' não foi encontrado.")